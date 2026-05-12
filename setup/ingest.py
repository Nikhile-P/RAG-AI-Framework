import os
import json
import csv
from dotenv import load_dotenv
from langchain_community.document_loaders import TextLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import Chroma
from langchain_core.documents import Document

load_dotenv()

# Absolute paths so this script works regardless of where it's called from
_BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
data_folder = os.path.join(_BASE, "documents")
_VECTOR_DIR = os.path.join(_BASE, "vector_db")

# 1. LOAD DATA WITH METADATA (The "Provenance" Layer)
documents = []

def add_text_docs(file_path, file_name):
    loader = TextLoader(file_path, encoding="utf-8")
    file_docs = loader.load()
    for doc in file_docs:
        doc.metadata["source"] = file_name
    documents.extend(file_docs)


def add_csv_docs(file_path, file_name):
    rows = []
    with open(file_path, "r", encoding="utf-8", newline="") as fh:
        reader = csv.DictReader(fh)
        for row in reader:
            rows.append(row)
    content = "\n".join(json.dumps(row, ensure_ascii=False) for row in rows)
    documents.append(Document(page_content=content, metadata={"source": file_name}))


def add_json_docs(file_path, file_name):
    with open(file_path, "r", encoding="utf-8") as fh:
        payload = json.load(fh)
    content = json.dumps(payload, ensure_ascii=False, indent=2)
    documents.append(Document(page_content=content, metadata={"source": file_name}))


def add_docx_docs(file_path, file_name):
    try:
        from docx import Document as DocxDocument
    except ImportError:
        print(f"Skipping {file_name}: python-docx not installed")
        return

    doc = DocxDocument(file_path)
    text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    if text.strip():
        documents.append(Document(page_content=text, metadata={"source": file_name}))


def add_pdf_docs(file_path, file_name):
    try:
        from pypdf import PdfReader
    except ImportError:
        print(f"Skipping {file_name}: pypdf not installed")
        return

    reader = PdfReader(file_path)
    pages = []
    for idx, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        if text:
            pages.append(f"[Page {idx}]\n{text}")
    if pages:
        documents.append(Document(page_content="\n\n".join(pages), metadata={"source": file_name}))


def add_xlsx_docs(file_path, file_name):
    try:
        from openpyxl import load_workbook
    except ImportError:
        print(f"Skipping {file_name}: openpyxl not installed")
        return

    wb = load_workbook(file_path, data_only=True)
    sheet_text = []
    for ws in wb.worksheets:
        sheet_text.append(f"[Sheet: {ws.title}]")
        rows = ws.iter_rows(values_only=True)
        for row in rows:
            values = [str(cell) for cell in row if cell is not None and str(cell).strip()]
            if values:
                sheet_text.append(" | ".join(values))
    text = "\n".join(sheet_text).strip()
    if text:
        documents.append(Document(page_content=text, metadata={"source": file_name}))


for file in os.listdir(data_folder):
    file_path = os.path.join(data_folder, file)
    lower = file.lower()

    if lower.endswith((".txt", ".md")):
        add_text_docs(file_path, file)
    elif lower.endswith(".csv"):
        add_csv_docs(file_path, file)
    elif lower.endswith(".json"):
        add_json_docs(file_path, file)
    elif lower.endswith(".docx"):
        add_docx_docs(file_path, file)
    elif lower.endswith(".pdf"):
        add_pdf_docs(file_path, file)
    elif lower.endswith(".xlsx"):
        add_xlsx_docs(file_path, file)

# 2. TRANSFORM DATA (The "Chunking" Layer)
# CS Tip: Overlap prevents context loss at the edges of a chunk
text_splitter = RecursiveCharacterTextSplitter(
    chunk_size=800,
    chunk_overlap=150
)
docs = text_splitter.split_documents(documents)

# 3. EMBED & STORE (The "Persistence" Layer)
embeddings = HuggingFaceEmbeddings(
    model_name="sentence-transformers/all-MiniLM-L6-v2"
)

# Store in ChromaDB
db = Chroma.from_documents(
    docs,
    embeddings,
    persist_directory=_VECTOR_DIR
)

print(f"Done — {len(docs)} chunks indexed into vector_db at {_VECTOR_DIR}")