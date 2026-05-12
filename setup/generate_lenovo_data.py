from pathlib import Path
import json
import csv
from datetime import date

DATA_DIR = Path("documents")
DATA_DIR.mkdir(parents=True, exist_ok=True)


def write_text_file(name: str, content: str) -> None:
    (DATA_DIR / name).write_text(content, encoding="utf-8")


# 1) TXT: Operations brief
write_text_file(
    "Lenovo_Global_Operations_Brief_2026.txt",
    """Lenovo Global Operations Brief 2026

Executive Context:
Lenovo enterprise operations focus on resilient supply chains, customer delivery SLAs, and high-performance infrastructure support.

Regional Priorities:
- North America: Expand enterprise server install base and managed deployment services.
- EMEA: Increase smart edge deployments for manufacturing and public sector.
- APAC: Strengthen hyperscale partnerships and sovereign cloud compliance pathways.

Operational KPIs:
- Average deployment lead time target: 14 business days.
- Field issue first-response target: 2 hours.
- Critical incident resolution target: 8 hours.

Risk Controls:
- Dual-sourcing for key components where feasible.
- Quarterly stress testing for logistics continuity.
- Formal firmware governance and staged rollout approvals.
""",
)

# 2) TXT: Product matrix
write_text_file(
    "Lenovo_Server_Portfolio_Matrix_2026.txt",
    """Lenovo Server Portfolio Matrix 2026

ThinkSystem SR670 V2:
- Use case: GPU-dense high-performance computing
- Typical sectors: Healthcare imaging, autonomous systems, industrial analytics
- Key value: Balanced performance-per-watt for mixed workloads

ThinkSystem SR650 V3:
- Use case: Enterprise virtualization and advanced analytics
- Typical sectors: Banking, insurance, retail analytics
- Key value: Flexible PCIe expansion with broad software compatibility

ThinkEdge SE450:
- Use case: Edge processing in constrained environments
- Typical sectors: Logistics, retail stores, smart manufacturing
- Key value: Ruggedized form factor with remote manageability
""",
)

# 3) CSV: Quarterly revenue mock
with (DATA_DIR / "Lenovo_Enterprise_Revenue_By_Quarter_2024_2026.csv").open("w", encoding="utf-8", newline="") as fh:
    writer = csv.DictWriter(
        fh,
        fieldnames=["year", "quarter", "region", "business_unit", "revenue_usd_m", "gross_margin_pct"],
    )
    writer.writeheader()
    rows = [
        (2024, "Q1", "NA", "Advanced Infrastructure", 420, 31.2),
        (2024, "Q2", "EMEA", "Advanced Infrastructure", 398, 29.8),
        (2024, "Q3", "APAC", "Advanced Infrastructure", 442, 30.4),
        (2024, "Q4", "Global", "Services", 515, 34.1),
        (2025, "Q1", "NA", "Advanced Infrastructure", 468, 32.0),
        (2025, "Q2", "EMEA", "Services", 536, 34.8),
        (2025, "Q3", "APAC", "Advanced Infrastructure", 489, 31.3),
        (2025, "Q4", "Global", "Services", 562, 35.2),
        (2026, "Q1", "NA", "Advanced Infrastructure", 505, 32.7),
        (2026, "Q2", "EMEA", "Advanced Infrastructure", 473, 31.1),
    ]
    for y, q, r, bu, rev, gm in rows:
        writer.writerow(
            {
                "year": y,
                "quarter": q,
                "region": r,
                "business_unit": bu,
                "revenue_usd_m": rev,
                "gross_margin_pct": gm,
            }
        )

# 4) CSV: Support SLA
with (DATA_DIR / "Lenovo_Support_SLA_Performance_2026.csv").open("w", encoding="utf-8", newline="") as fh:
    writer = csv.DictWriter(
        fh,
        fieldnames=["month", "priority", "tickets", "first_response_minutes_avg", "resolution_hours_avg", "sla_met_pct"],
    )
    writer.writeheader()
    for month, p1, p2 in [
        ("Jan", (122, 18, 7.1, 97.6), (301, 64, 18.2, 94.5)),
        ("Feb", (115, 16, 6.8, 98.1), (288, 61, 17.6, 95.1)),
        ("Mar", (130, 19, 7.4, 97.2), (326, 66, 19.0, 93.9)),
        ("Apr", (126, 17, 7.0, 97.8), (314, 62, 18.4, 94.7)),
    ]:
        writer.writerow({"month": month, "priority": "P1", "tickets": p1[0], "first_response_minutes_avg": p1[1], "resolution_hours_avg": p1[2], "sla_met_pct": p1[3]})
        writer.writerow({"month": month, "priority": "P2", "tickets": p2[0], "first_response_minutes_avg": p2[1], "resolution_hours_avg": p2[2], "sla_met_pct": p2[3]})

# 5) JSON: Deal pipeline
pipeline = {
    "report_date": str(date.today()),
    "currency": "USD",
    "opportunities": [
        {"account": "Apex Health Network", "region": "NA", "segment": "Healthcare", "stage": "Proposal", "value_m": 18.5, "solution": "SR670 V2 Cluster + Managed Operations"},
        {"account": "EuroFab Group", "region": "EMEA", "segment": "Manufacturing", "stage": "Negotiation", "value_m": 11.2, "solution": "Smart Edge + Factory Vision"},
        {"account": "PacificTel", "region": "APAC", "segment": "Telecom", "stage": "Discovery", "value_m": 7.8, "solution": "Processing Grid for 5G Analytics"},
    ],
}
(DATA_DIR / "Lenovo_Enterprise_Deal_Pipeline_2026.json").write_text(json.dumps(pipeline, indent=2), encoding="utf-8")

# 6) JSON: Compliance controls
compliance = {
    "framework": "Internal Systems Infrastructure Governance",
    "controls": [
        {"id": "SEC-001", "name": "System artifact integrity", "status": "Implemented"},
        {"id": "OPS-014", "name": "GPU firmware patch window", "status": "Implemented"},
        {"id": "RISK-022", "name": "Vendor continuity review", "status": "In Progress"},
    ],
}
(DATA_DIR / "Lenovo_Governance_Controls_2026.json").write_text(json.dumps(compliance, indent=2), encoding="utf-8")

# 7) Markdown: Business strategy note
write_text_file(
    "Lenovo_Business_Strategy_Notebook_2026.md",
    """# Lenovo Business Strategy Notebook 2026

## Strategic Themes
- Accelerate high-performance infrastructure wins in regulated sectors.
- Bundle services to improve customer lifetime value.
- Build repeatable deployment playbooks for enterprise solutions.

## Competitive Positioning
Lenovo emphasizes deployment reliability, platform openness, and enterprise-grade service coverage.

## Executive Action List
1. Expand strategic channel partnerships for solution delivery.
2. Standardize technical readiness assessments in pre-sales cycle.
3. Track post-deployment ROI outcomes at 30/90/180 day windows.
""",
)

# 8) Markdown: Product brief
write_text_file(
    "Lenovo_Product_Engineering_Notes_SR670V2.md",
    """# SR670 V2 Engineering Notes

- Thermal design validated for sustained high-density GPU workloads.
- BIOS tuning profile optimized for mixed high-throughput jobs.
- Supports enterprise observability integrations for GPU fleet health.

## Validation Checklist
- Power capping scenarios
- PCIe throughput saturation tests
- Failover and restart consistency
""",
)

# 9) XLSX
try:
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = "CapacityPlan"
    ws.append(["Site", "RackCount", "GPU_Nodes", "UtilizationPct", "ExpansionQuarter"])
    ws.append(["Chicago", 24, 96, 73, "Q3-2026"])
    ws.append(["Frankfurt", 18, 72, 69, "Q4-2026"])
    ws.append(["Singapore", 20, 80, 77, "Q1-2027"])

    ws2 = wb.create_sheet("Budget")
    ws2.append(["CostCenter", "CapexUSDm", "OpexUSDm", "Approved"])
    ws2.append(["Advanced-Infrastructure", 12.5, 4.2, "Yes"])
    ws2.append(["Managed-Services", 4.0, 2.8, "Yes"])
    ws2.append(["Training-Programs", 1.2, 0.9, "Pending"])

    wb.save(DATA_DIR / "Lenovo_Capacity_and_Budget_2026.xlsx")
except Exception as ex:
    print(f"Skipping xlsx generation: {ex}")

# 10) DOCX
try:
    from docx import Document

    doc = Document()
    doc.add_heading("Lenovo Enterprise Systems Implementation Guide", level=1)
    doc.add_paragraph("This internal guide summarizes deployment phases, ownership responsibilities, and governance checkpoints.")
    doc.add_heading("Deployment Phases", level=2)
    doc.add_paragraph("Phase 1: Assessment and architecture validation")
    doc.add_paragraph("Phase 2: Pilot rollout and baseline KPI measurement")
    doc.add_paragraph("Phase 3: Production hardening and operational handover")
    doc.add_heading("KPI Baseline", level=2)
    doc.add_paragraph("Track latency, throughput, and incident recovery SLA weekly.")
    doc.save(DATA_DIR / "Lenovo_Enterprise_Implementation_Guide.docx")
except Exception as ex:
    print(f"Skipping docx generation: {ex}")

# 11) PDF
try:
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas

    pdf_path = DATA_DIR / "Lenovo_Quarterly_Program_Review_Q2_2026.pdf"
    c = canvas.Canvas(str(pdf_path), pagesize=letter)
    width, height = letter
    y = height - 72

    c.setFont("Helvetica-Bold", 16)
    c.drawString(72, y, "Lenovo Quarterly Program Review")
    y -= 30

    c.setFont("Helvetica", 11)
    lines = [
        "Program Highlights:",
        "- Increased advanced infrastructure adoption in healthcare and manufacturing.",
        "- Improved average deployment lead time through standardized runbooks.",
        "- Expanded managed services attach rate across strategic accounts.",
        "",
        "Key Risks:",
        "- Component lead-time volatility in select regions.",
        "- Skills gap in specialized infrastructure operational teams.",
        "",
        "Mitigations:",
        "- Buffer inventory for critical GPU and networking components.",
        "- Expanded internal enablement and certification pathways.",
    ]
    for line in lines:
        c.drawString(72, y, line)
        y -= 16
        if y < 72:
            c.showPage()
            c.setFont("Helvetica", 11)
            y = height - 72

    c.save()
except Exception as ex:
    print(f"Skipping pdf generation: {ex}")

# 12) TXT: procurement insights
write_text_file(
    "Lenovo_Procurement_and_Pricing_Intelligence_2026.txt",
    """Lenovo Procurement and Pricing Intelligence 2026

Pricing Levers:
- Volume-based discounts for multi-rack deployments.
- Service bundle incentives tied to multi-year agreements.
- Preferred pricing windows aligned with quarterly planning cycles.

Procurement Best Practices:
- Lock strategic components 1-2 quarters ahead.
- Align legal and compliance approvals in pre-proposal phase.
- Use standard SoW templates for faster contracting.
""",
)

print("Generated Lenovo internal corpus files in documents/")
